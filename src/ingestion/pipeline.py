"""Multi-format document ingestion. PDF (pdfplumber/pypdf), TXT, CSV, JSON, DOCX, XLSX."""
from __future__ import annotations

import hashlib
import json
import logging
import uuid
from pathlib import Path

from src.abstractions.ledger import LedgerEventType
from src.schemas.models import IngestedDocument

logger = logging.getLogger(__name__)
MAX_FILE_SIZE_MB = 50
MAX_PAGES = 2000

# Extensions we can ingest and extract text from (lowercase for comparison)
SUPPORTED_EXTENSIONS = {".pdf", ".txt", ".csv", ".json", ".docx", ".xlsx", ".xls"}
CONTENT_TYPES = {
    ".pdf": "application/pdf",
    ".txt": "text/plain",
    ".csv": "text/csv",
    ".json": "application/json",
    ".docx": "application/vnd.openxmlformats-officedocument.wordprocessingml.document",
    ".xlsx": "application/vnd.openxmlformats-officedocument.spreadsheetml.sheet",
    ".xls": "application/vnd.ms-excel",
}


def _compute_sha256(path: Path) -> str:
    h = hashlib.sha256()
    with path.open("rb") as f:
        for chunk in iter(lambda: f.read(65536), b""):
            h.update(chunk)
    return h.hexdigest()


def _extract_text_pypdf(path: Path) -> tuple[str, int]:
    from pypdf import PdfReader
    reader = PdfReader(str(path))
    pages = reader.pages
    n = len(pages)
    if n > MAX_PAGES:
        raise ValueError("PDF too many pages")
    parts = []
    for i, page in enumerate(pages):
        if i >= 500:
            break
        try:
            parts.append(page.extract_text() or "")
        except Exception as e:
            logger.warning("Page %s failed: %s", i + 1, e)
    return "\n\n".join(parts), n


def _extract_text_pdfplumber(path: Path) -> tuple[str, int]:
    import pdfplumber
    parts = []
    with pdfplumber.open(path) as pdf:
        n = len(pdf.pages)
        if n > MAX_PAGES:
            raise ValueError("PDF too many pages")
        for i, page in enumerate(pdf.pages):
            if i >= 500:
                break
            try:
                text = page.extract_text() or ""
                tables = page.extract_tables()
                if tables:
                    for table in tables:
                        for row in table:
                            text += " " + " ".join(str(c) if c else "" for c in row)
                parts.append(text)
            except Exception as e:
                logger.warning("Page %s failed: %s", i + 1, e)
    return "\n\n".join(parts), n


def _extract_text_pdf(path: Path) -> tuple[str, int]:
    try:
        return _extract_text_pdfplumber(path)
    except Exception as e:
        logger.debug("pdfplumber failed, using pypdf: %s", e)
        return _extract_text_pypdf(path)


def _extract_text_txt(path: Path) -> tuple[str, int]:
    raw = path.read_bytes()
    for enc in ("utf-8", "utf-8-sig", "latin-1", "cp1252"):
        try:
            text = raw.decode(enc)
            return (text.strip() or "", 1)
        except UnicodeDecodeError:
            continue
    return (raw.decode("utf-8", errors="replace"), 1)


def _extract_text_csv(path: Path) -> tuple[str, int]:
    text, _ = _extract_text_txt(path)
    return (text, 1)


def _extract_text_json(path: Path) -> tuple[str, int]:
    raw = path.read_bytes()
    try:
        data = json.loads(raw.decode("utf-8"))
        text = json.dumps(data, indent=0, ensure_ascii=False)
    except Exception:
        text = raw.decode("utf-8", errors="replace")
    return (text, 1)


def _extract_text_docx(path: Path) -> tuple[str, int]:
    from docx import Document
    doc = Document(str(path))
    parts = []
    for para in doc.paragraphs:
        if para.text.strip():
            parts.append(para.text)
    for table in doc.tables:
        for row in table.rows:
            parts.append(" ".join(str(c.text or "").strip() for c in row.cells))
    text = "\n".join(parts)
    page_approx = max(1, (len(parts) // 30) + 1)
    return (text, page_approx)


def _extract_text_xlsx(path: Path) -> tuple[str, int]:
    import openpyxl
    wb = openpyxl.load_workbook(path, read_only=True, data_only=True)
    parts = []
    sheet_count = 0
    for sheet in wb.worksheets:
        sheet_count += 1
        for row in sheet.iter_rows(values_only=True):
            parts.append(" ".join(str(c) if c is not None else "" for c in row))
    wb.close()
    text = "\n".join(parts)
    return (text, max(1, sheet_count))


def _extract_text(path: Path, ext: str) -> tuple[str, int]:
    ext_lower = ext.lower()
    if ext_lower == ".pdf":
        return _extract_text_pdf(path)
    if ext_lower == ".txt":
        return _extract_text_txt(path)
    if ext_lower == ".csv":
        return _extract_text_csv(path)
    if ext_lower == ".json":
        return _extract_text_json(path)
    if ext_lower == ".docx":
        return _extract_text_docx(path)
    if ext_lower in (".xlsx", ".xls"):
        return _extract_text_xlsx(path)
    raise ValueError(f"Unsupported extension: {ext}")


def _iter_supported_files(docs_path: Path):
    for p in docs_path.rglob("*"):
        if not p.is_file():
            continue
        if p.suffix.lower() in SUPPORTED_EXTENSIONS:
            yield p


def ingest_documents(case_id: str, docs_dir: str | Path, storage, database, ledger=None):
    docs_path = Path(docs_dir)
    if not docs_path.is_dir():
        raise FileNotFoundError("Not a directory: " + str(docs_dir))
    max_bytes = MAX_FILE_SIZE_MB * 1024 * 1024
    results = []
    for path in _iter_supported_files(docs_path):
        if path.stat().st_size > max_bytes or path.stat().st_size == 0:
            continue
        ext = path.suffix.lower()
        content_type = CONTENT_TYPES.get(ext, "application/octet-stream")
        content_hash = _compute_sha256(path)
        try:
            text_content, page_count = _extract_text(path, ext)
            logger.info("Extracted %d chars from %s", len(text_content or ""), path.name)
            if not text_content.strip() and ext == ".pdf":
                logger.warning("Empty text extracted from %s - likely a scanned image/PDF requiring OCR.", path.name)
        except Exception as e:
            logger.warning("Extract failed %s: %s", path.name, e)
            continue
        doc_id = str(uuid.uuid4()).replace("-", "")[:24]
        storage_key = "cases/" + case_id + "/docs/" + doc_id + "_" + path.name
        storage.put(storage_key, path.read_bytes(), content_type=content_type)
        text_preview = (text_content or "")[:10000]
        database.insert_document(case_id, doc_id, path.name, content_hash, storage_key, page_count, None)
        doc = IngestedDocument(
            doc_id=doc_id,
            case_id=case_id,
            filename=path.name,
            content_hash=content_hash,
            storage_key=storage_key,
            page_count=page_count,
            text_preview=text_preview,
            metadata={},
        )
        results.append(doc)
        if ledger:
            ledger.append(
                LedgerEventType.DOCUMENT_INGESTED.value,
                {"doc_id": doc_id, "case_id": case_id, "filename": path.name, "content_hash": content_hash},
                entity_id=doc_id,
            )
        logger.info("Ingested %s -> %s", path.name, doc_id)
    return results
